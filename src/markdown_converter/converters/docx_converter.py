from pathlib import Path
from docx import Document
from markdown_converter.config import ConvertOptions
from markdown_converter.converters.base import BaseConverter
from markdown_converter.utils import normalize_text, table_to_markdown


class DocxConverter(BaseConverter):
    def convert(self, input_path: Path, options: ConvertOptions) -> str:
        doc = Document(str(input_path))
        parts: list[str] = []

        for block in self._iter_blocks(doc):
            kind, value = block
            if kind == "paragraph":
                md = self._paragraph_to_markdown(value)
                if md:
                    parts.append(md)
            elif kind == "table":
                table_md = table_to_markdown(value)
                if table_md:
                    parts.append(table_md)

        return normalize_text("\n\n".join(parts))

    def _iter_blocks(self, doc: Document):
        for paragraph in doc.paragraphs:
            yield "paragraph", paragraph
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            yield "table", rows

    def _paragraph_to_markdown(self, paragraph) -> str:
        text = paragraph.text.strip()
        if not text:
            return ""

        style_name = (paragraph.style.name or "").lower()

        if "heading 1" in style_name:
            return f"# {text}"
        if "heading 2" in style_name:
            return f"## {text}"
        if "heading 3" in style_name:
            return f"### {text}"
        if "heading 4" in style_name:
            return f"#### {text}"
        if "list bullet" in style_name:
            return f"- {text}"
        if "list number" in style_name:
            return f"1. {text}"

        chunks: list[str] = []
        for run in paragraph.runs:
            run_text = run.text
            if not run_text:
                continue
            if run.bold and run.italic:
                run_text = f"***{run_text}***"
            elif run.bold:
                run_text = f"**{run_text}**"
            elif run.italic:
                run_text = f"*{run_text}*"
            chunks.append(run_text)

        return "".join(chunks).strip() or text
